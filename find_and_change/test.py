"""
教程地址
https://zhuanlan.zhihu.com/p/90855359
"""
import docx


def print_all_doc():
    # 每一段的内容
    for para in doc.paragraphs:
        print(para.text)
    # 每一段的编号、内容
    for i in range(len(doc.paragraphs)):
        print(str(i), doc.paragraphs[i].text)


def check_and_change(document, replace_dict):
    """
    遍历word中的所有 paragraphs，在每一段中发现含有key 的内容，就替换为 value 。
   （key 和 value 都是replace_dict中的键值对。）
    """
    for para in document.paragraphs:
        for i in range(len(para.runs)):
            for key, value in replace_dict.items():
                if key in para.runs[i].text:
                    print(key + "->" + value)
                    para.runs[i].text = para.runs[i].text.replace(key, value)
    return document


if __name__ == '__main__':
    doc = docx.Document("共产党宣言origin.docx")
    #  将想要替换的内容写成字典的形式，
    #  dict = {"想要被替换的字符串": "新的字符串"}
    replace_dict = {
        "共产党": "the Communist Party",
        "？": "?",

    }
    new_doc = check_and_change(doc, replace_dict)
    new_doc.save("共产党宣言replaced.docx")
    print_all_doc()
